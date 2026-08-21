import html
import os
import re
from pathlib import Path

import fitz
from docx import Document
from docx.shared import Mm
from pydantic import BaseModel, Field
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.models import ApplicationDraft, CandidateProfile, JobOffer, MatchResult
from app.services.application_drafts import verified_skills_for_draft
from app.services.technology_matcher import normalize


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_OUTPUT_ROOT = PROJECT_ROOT / "generated_documents"
PLACEHOLDER_PATTERNS = (
    "a completer",
    "à compléter",
    "todo",
    "undefined",
    "null",
    "{{",
    "}}",
)


class DocumentValidationReport(BaseModel):
    valid: bool
    errors: list[str] = Field(default_factory=list)
    extracted_files: dict[str, str] = Field(default_factory=dict)


class GeneratedApplicationDocuments(BaseModel):
    draft_id: int
    version: int
    cover_letter_docx: str
    cover_letter_pdf: str
    adapted_cv_pdf: str
    validation: DocumentValidationReport


def output_root() -> Path:
    configured = os.getenv("GENERATED_DOCUMENTS_DIR", "").strip()
    return Path(configured).resolve() if configured else DEFAULT_OUTPUT_ROOT


def draft_directory(draft: ApplicationDraft) -> Path:
    path = output_root() / f"draft-{draft.id}" / f"version-{draft.version}"
    path.mkdir(parents=True, exist_ok=True)
    return path


def write_docx(text: str, destination: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

    for block in text.split("\n\n"):
        paragraph = document.add_paragraph()
        paragraph.add_run(block.strip())
        paragraph.paragraph_format.space_after = Mm(8)

    document.save(destination)


def write_reportlab_pdf(
    text: str,
    destination: Path,
    *,
    title: str,
) -> None:
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "ApplyMatchBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        alignment=TA_LEFT,
        spaceAfter=7 * mm,
    )
    title_style = ParagraphStyle(
        "ApplyMatchTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=18,
        spaceAfter=10 * mm,
    )
    story = [Paragraph(html.escape(title), title_style)]

    for block in text.split("\n\n"):
        cleaned = block.strip()
        if not cleaned:
            continue
        story.append(
            Paragraph(
                html.escape(cleaned).replace("\n", "<br/>") ,
                body_style,
            )
        )
        story.append(Spacer(1, 2 * mm))

    document = SimpleDocTemplate(
        str(destination),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title,
    )
    document.build(story)


def write_weasyprint_pdf(
    text: str,
    destination: Path,
    *,
    title: str,
) -> None:
    from weasyprint import HTML

    paragraphs = "".join(
        f"<p>{html.escape(block).replace(chr(10), '<br>')}</p>"
        for block in text.split("\n\n")
        if block.strip()
    )
    content = f"""
    <!doctype html><html lang="fr"><head><meta charset="utf-8">
    <style>@page {{ size: A4; margin: 20mm; }}
    body {{ font-family: Arial, sans-serif; font-size: 10.5pt;
    line-height: 1.45; color: #111827; }} h1 {{ font-size: 15pt; }}</style>
    </head><body><h1>{html.escape(title)}</h1>{paragraphs}</body></html>
    """
    HTML(string=content).write_pdf(destination)


def write_pdf(
    text: str,
    destination: Path,
    *,
    title: str,
) -> None:
    if os.getenv("PDF_ENGINE", "reportlab").strip().casefold() == "weasyprint":
        try:
            write_weasyprint_pdf(text, destination, title=title)
            return
        except Exception:
            pass

    write_reportlab_pdf(text, destination, title=title)


def extract_pdf_text(path: Path) -> str:
    with fitz.open(path) as document:
        return "\n".join(page.get_text() for page in document)


def competency_section(text: str) -> str:
    normalized_text = normalize(text)
    match = re.search(
        r"competences pertinentes\s+(.*?)(?:\s+experiences|\s+projets|"
        r"\s+disponibilite|$)",
        normalized_text,
        re.DOTALL,
    )
    return match.group(1) if match else ""


def validate_documents(
    *,
    letter_pdf: Path,
    cv_pdf: Path,
    profile: CandidateProfile,
    offer: JobOffer,
    match_result: MatchResult,
) -> DocumentValidationReport:
    extracted = {
        "cover_letter_pdf": extract_pdf_text(letter_pdf),
        "adapted_cv_pdf": extract_pdf_text(cv_pdf),
    }
    letter = normalize(extracted["cover_letter_pdf"])
    cv = normalize(extracted["adapted_cv_pdf"])
    combined = f"{letter}\n{cv}"
    errors: list[str] = []

    required_values = {
        "nom": profile.full_name,
        "entreprise": offer.company,
        "poste": offer.title,
    }
    for label, value in required_values.items():
        target_text = combined if label == "nom" else letter
        if normalize(value) not in target_text:
            errors.append(f"champ_manquant:{label}")

    for pattern in PLACEHOLDER_PATTERNS:
        if normalize(pattern) in combined:
            errors.append(f"champ_vide:{pattern}")

    verified = set(verified_skills_for_draft(match_result))
    cv_skills = competency_section(extracted["adapted_cv_pdf"])
    for technology in match_result.unknown_technologies:
        if normalize(technology) in cv_skills:
            errors.append(f"technologie_non_prouvee:{technology}")

    for technology in verified:
        if normalize(technology) not in cv_skills:
            errors.append(f"technologie_connue_absente:{technology}")

    return DocumentValidationReport(
        valid=not errors,
        errors=list(dict.fromkeys(errors)),
        extracted_files=extracted,
    )


def generate_application_documents(
    *,
    draft: ApplicationDraft,
    profile: CandidateProfile,
    offer: JobOffer,
    match_result: MatchResult,
) -> GeneratedApplicationDocuments:
    directory = draft_directory(draft)
    letter_docx = directory / "lettre-motivation.docx"
    letter_pdf = directory / "lettre-motivation.pdf"
    cv_pdf = directory / "cv-adapte.pdf"

    write_docx(draft.cover_letter, letter_docx)
    write_pdf(
        draft.cover_letter,
        letter_pdf,
        title=f"Lettre de motivation – {offer.company}",
    )
    write_pdf(
        draft.adapted_cv_snapshot,
        cv_pdf,
        title=f"CV adapté – {offer.title}",
    )
    validation = validate_documents(
        letter_pdf=letter_pdf,
        cv_pdf=cv_pdf,
        profile=profile,
        offer=offer,
        match_result=match_result,
    )

    return GeneratedApplicationDocuments(
        draft_id=draft.id,
        version=draft.version,
        cover_letter_docx=letter_docx.name,
        cover_letter_pdf=letter_pdf.name,
        adapted_cv_pdf=cv_pdf.name,
        validation=validation,
    )
